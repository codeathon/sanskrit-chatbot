#!/usr/bin/env python3
"""
VedaGPT — Data Ingestion Script
Scans ./data/ for supported files, builds knowledge base and search index.

Usage:
    python3 ingest.py

Supported in ./data/ (any names, any count):
  .docx  — full text chunked for search
  .xlsx  — if a sheet named "RV Database" exists, rows are ingested as mantra docs.
           Columns B, C, D (Excel) = mandala, sukta, verse; a composite index key
           "B.C.D" (e.g. 1.1.1) is stored as composite_key and used in doc ids and text.
           Other sheets are
           turned into plain text rows and chunked, except names listed in
           VEDAGPT_SKIP_XLSX_SHEETS (default skips tab "Website Design", which is
           not corpus text — it lives inside the .xlsx, not a separate .docx file).
Other file types are skipped with a short notice.
"""

import os
import re
import json
import math
import pickle
import shutil
import subprocess
from pathlib import Path
from collections import Counter, defaultdict

from openpyxl import load_workbook

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, 'data')
KB_DIR = os.path.join(BASE_DIR, 'knowledge-base')
os.makedirs(KB_DIR, exist_ok=True)
os.makedirs(DATA_DIR, exist_ok=True)


def generic_xlsx_sheets_to_skip():
    """
    Tab names on .xlsx files to omit from generic text ingest (never applies to
    'RV Database'). Unset env → skip 'Website Design'. Set VEDAGPT_SKIP_XLSX_SHEETS
    to empty to skip nothing; comma-separated list for multiple names (case-insensitive).
    """
    if 'VEDAGPT_SKIP_XLSX_SHEETS' in os.environ:
        raw = os.environ['VEDAGPT_SKIP_XLSX_SHEETS']
    else:
        raw = 'Website Design'
    return frozenset(x.strip().lower() for x in raw.split(',') if x.strip())


STOPWORDS = {
    'the', 'a', 'an', 'and', 'or', 'of', 'in', 'to', 'is', 'are', 'was', 'were', 'it', 'this', 'that',
    'with', 'for', 'on', 'at', 'by', 'from', 'as', 'be', 'been', 'being', 'have', 'has', 'had',
    'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'not', 'but', 'we',
    'they', 'he', 'she', 'i', 'you', 'his', 'her', 'their', 'our', 'its', 'which', 'who', 'what',
    'when', 'where', 'how', 'all', 'also', 'more', 'into', 'than', 'them', 'these', 'those',
    'so', 'if', 'about', 'up', 'out', 'no', 'can', 'my', 'your', 'there', 'then', 'only', 'after'
}


def slug_label(s):
    """Stable id fragment from a file or sheet name."""
    s = re.sub(r'[^a-zA-Z0-9]+', '_', str(s)).strip('_')
    return s or 'source'


def list_ingestible_files():
    """Return sorted paths under DATA_DIR for .docx / .xlsx only."""
    if not os.path.isdir(DATA_DIR):
        return []
    out = []
    for name in sorted(os.listdir(DATA_DIR)):
        if name.startswith('.'):
            continue
        path = os.path.join(DATA_DIR, name)
        if not os.path.isfile(path):
            continue
        low = name.lower()
        if low.endswith('.docx') or low.endswith('.xlsx'):
            out.append(path)
        else:
            print(f"  (skip unsupported: {name})")
    return out


def _extract_docx_via_python_docx(path):
    # Pure-Python path when pandoc is not installed (see extract_docx).
    from docx import Document
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(' '.join(cells))
    return '\n\n'.join(parts)


def extract_docx(path):
    text = None
    pandoc_bin = shutil.which('pandoc')
    if pandoc_bin:
        result = subprocess.run(
            [pandoc_bin, path, '-t', 'plain'], capture_output=True, text=True
        )
        if result.returncode == 0 and result.stdout.strip():
            text = result.stdout
    if not text:
        text = _extract_docx_via_python_docx(path)
    text = re.sub(r'\nPage \d+\n', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def chunk_text(text, source, chunk_size=1200, overlap=200):
    chunks = []
    words = text.split()
    i = 0
    cid = 0
    while i < len(words):
        chunk = ' '.join(words[i:i + chunk_size])
        chunks.append({'id': f"{source}_{cid}", 'source': source, 'text': chunk, 'type': 'text'})
        cid += 1
        i += chunk_size - overlap
    return chunks


def tokenize(text):
    tokens = re.findall(r'\b[a-zA-Zāīūṛḷṃḥśṣṭḍṇñṅ]+\b', text.lower())
    return [t for t in tokens if t not in STOPWORDS and len(t) > 2]


def build_tfidf(all_docs):
    print("  Tokenizing documents...")
    doc_tokens = [tokenize(d['text']) for d in all_docs]
    n = len(all_docs)
    df = defaultdict(int)
    for tokens in doc_tokens:
        for tok in set(tokens):
            df[tok] += 1
    idf = {tok: math.log(n / (1 + df[tok])) for tok in df}
    print(f"  Vocabulary: {len(idf)} terms")
    tfidf_vecs = []
    for tokens in doc_tokens:
        tf = Counter(tokens)
        total = len(tokens) if tokens else 1
        vec = {tok: (cnt / total) * idf[tok] for tok, cnt in tf.items() if tok in idf}
        tfidf_vecs.append(vec)
    return tfidf_vecs, idf


def ingest_docx_path(path, all_docs, source_meta):
    fname = os.path.basename(path)
    base = slug_label(Path(path).stem)
    print(f"\nExtracting {fname} (.docx)...")
    try:
        text = extract_docx(path)
    except Exception as e:
        print(f"  ⚠️  Failed to read {fname}: {e}")
        return
    if not text:
        print("  → (empty after extract, skipped)")
        return
    chunks = chunk_text(text, base)
    all_docs.extend(chunks)
    source_meta.append({
        'name': base,
        'description': Path(path).stem.replace('_', ' '),
        'chunks': len(chunks)
    })
    print(f"  → {len(chunks)} chunks, {len(text):,} chars")


def sheet_to_plain_text(ws):
    """Join non-empty cells per row; one row per line (generic spreadsheet)."""
    lines = []
    for row in ws.iter_rows(values_only=True):
        cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
        if cells:
            lines.append(' | '.join(cells))
    return '\n'.join(lines).strip()


def _xlsx_bcd_part(cell):
    """
    Normalize one Excel cell for mandala / sukta / verse (columns B, C, D).
    Int-like floats from spreadsheets become string integers (e.g. 1.0 → '1').
    """
    if cell is None:
        return ''
    if isinstance(cell, float):
        if cell != cell:  # NaN
            return ''
        if cell == int(cell):
            return str(int(cell))
    return str(cell).strip()


def _composite_key_bcd(row):
    """Build composite index key from columns B, C, D (indices 1,2,3), e.g. '1.1.1'."""
    b, c, d = _xlsx_bcd_part(row[1]), _xlsx_bcd_part(row[2]), _xlsx_bcd_part(row[3])
    if not (b and c and d):
        return ''
    return f'{b}.{c}.{d}'


def ingest_rv_database_sheet(ws, file_base, fname, all_docs, source_meta):
    """
    Legacy VedaGPT blueprint layout: column A mantra #, B–D mandala/sukta/verse,
    Devanagari in column M (index 12). composite_key = B.C.D for stable indexing.
    """
    rows = list(ws.iter_rows(values_only=True))
    if len(rows) < 2:
        return 0
    src = f"{file_base}_Rig_Veda"
    rv_mantras = []
    for row in rows[1:]:
        if row[0] is None or not row[12]:
            continue
        mandala = _xlsx_bcd_part(row[1])
        sukta = _xlsx_bcd_part(row[2])
        verse = _xlsx_bcd_part(row[3])
        composite_key = _composite_key_bcd(row)
        mantra_devanagari = str(row[12])
        devata = str(row[9]) if row[9] else ''
        transliteration = str(row[14]) if row[14] else ''
        md = {
            'mantra_number': row[0],
            # Columns B, C, D (normalized strings for search and display)
            'mandala': mandala,
            'sukta': sukta,
            'verse': verse,
            'composite_key': composite_key,
            'devata': devata, 'chhandas': str(row[10]) if row[10] else '',
            'mantra_devanagari': mantra_devanagari,
            'pada_patha': str(row[13]) if row[13] else '',
            'transliteration': transliteration,
            'page_vol15': str(row[5]) if row[5] else '',
            'page_vol14': str(row[6]) if row[6] else '',
            'page_vol16': str(row[7]) if row[7] else '',
        }
        # Reference line uses composite B.C.D when present (same as a virtual index column).
        ref = composite_key if composite_key else f'{row[1]}.{row[2]}.{row[3]}'
        text = (f"Rig Veda {ref}: {mantra_devanagari} "
                f"Devata: {devata} Transliteration: {transliteration}")
        # Document id keyed by composite B.C.D; fallback to mantra_number (column A)
        if composite_key:
            rid = f"{src}_{composite_key}" #.replace('.', '_')}"
        else:
            rid = f"{src}_{row[0]}"
        doc = {'id': rid, 'source': src, 'type': 'mantra', 'text': text, 'mantra_data': md}
        rv_mantras.append(doc)
    if not rv_mantras:
        print(f"  → sheet 'RV Database' in {fname}: no mantra rows found (skipped)")
        return 0
    all_docs.extend(rv_mantras)
    source_meta.append({
        'name': src,
        'description': f'Rig Veda database — {fname}',
        'mantras': len(rv_mantras)
    })
    print(f"  → {fname} [RV Database]: {len(rv_mantras)} mantras")
    return len(rv_mantras)


def ingest_generic_xlsx_sheet(ws, file_base, sheet_name, fname, all_docs, source_meta):
    text = sheet_to_plain_text(ws)
    if not text:
        return
    src = f"{file_base}_{slug_label(sheet_name)}"
    chunks = chunk_text(text, src)
    all_docs.extend(chunks)
    source_meta.append({
        'name': src,
        'description': f'{sheet_name} — {fname}',
        'chunks': len(chunks)
    })
    print(f"  → {fname} [{sheet_name}]: {len(chunks)} text chunks")


def ingest_xlsx_path(path, all_docs, source_meta):
    fname = os.path.basename(path)
    base = slug_label(Path(path).stem)
    print(f"\nExtracting {fname} (.xlsx)...")
    try:
        wb = load_workbook(path, read_only=True)
    except Exception as e:
        hint = ''
        err = str(e).lower()
        if 'content_types' in err or 'archive' in err or 'not a zip' in err:
            hint = ' — file may be corrupt or not a valid .xlsx (Excel zip)'
        print(f"  ⚠️  Failed to open {fname}: {e}{hint}")
        return
    try:
        if 'RV Database' in wb.sheetnames:
            ingest_rv_database_sheet(wb['RV Database'], base, fname, all_docs, source_meta)

        skip_tabs = generic_xlsx_sheets_to_skip()
        for sheet_name in sorted(wb.sheetnames):
            if sheet_name == 'RV Database':
                continue
            if sheet_name.strip().lower() in skip_tabs:
                print(f"  → {fname} [{sheet_name}]: skipped (non-corpus sheet; see VEDAGPT_SKIP_XLSX_SHEETS)")
                continue
            try:
                ingest_generic_xlsx_sheet(
                    wb[sheet_name], base, sheet_name, fname, all_docs, source_meta
                )
            except Exception as e:
                print(f"  ⚠️  Sheet {sheet_name!r} in {fname}: {e}")
    finally:
        wb.close()


def main():
    print("=" * 60)
    print("  VedaGPT — Data Ingestion")
    print("=" * 60)
    print(f"  Scanning: {DATA_DIR}")

    all_docs = []
    source_meta = []
    paths = list_ingestible_files()

    if not paths:
        print("\n⚠️  No .docx or .xlsx files found in data/. Add files and run again.")
        return

    for path in paths:
        ext = Path(path).suffix.lower()
        if ext == '.docx':
            ingest_docx_path(path, all_docs, source_meta)
        elif ext == '.xlsx':
            ingest_xlsx_path(path, all_docs, source_meta)

    if not all_docs:
        print("\n⚠️  Nothing was indexed (empty or unreadable files).")
        return

    total_mantras = sum(1 for d in all_docs if d['type'] == 'mantra')
    print(f"\nTotal documents: {len(all_docs)} ({total_mantras} mantra records)")

    print("\nBuilding TF-IDF search index...")
    tfidf_vecs, idf = build_tfidf(all_docs)

    print("\nSaving knowledge base...")
    kb = {
        'metadata': {
            'sources': source_meta,
            'total_text_chunks': sum(s.get('chunks', 0) for s in source_meta),
            'total_rv_mantras': total_mantras
        },
        'text_chunks': [d for d in all_docs if d['type'] == 'text'],
        'rv_mantras': [d['mantra_data'] for d in all_docs if d['type'] == 'mantra']
    }
    with open(os.path.join(KB_DIR, 'knowledge_base.json'), 'w', encoding='utf-8') as f:
        json.dump(kb, f, ensure_ascii=False)

    index = {'all_docs': all_docs, 'tfidf_vecs': tfidf_vecs, 'idf': idf}
    with open(os.path.join(KB_DIR, 'search_index.pkl'), 'wb') as f:
        pickle.dump(index, f)

    size = os.path.getsize(os.path.join(KB_DIR, 'search_index.pkl')) / (1024 * 1024)
    print(f"  search_index.pkl — {size:.1f} MB")
    print("  knowledge_base.json — saved")
    print("\n✅ Ingestion complete! Run: python3 src/server/server.py")


if __name__ == '__main__':
    main()
